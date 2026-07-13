"""
Transcribir video (incluso de más de 1 hora) a texto, 100% gratis y local.

Motor de transcripción: faster-whisper (reimplementación optimizada de
Whisper con CTranslate2). Corre en tu CPU/GPU, sin necesidad de API key
ni conexión a internet una vez descargado el modelo.

INSTALACIÓN (una sola vez):
    pip install faster-whisper python-docx tqdm

    Además necesitas ffmpeg instalado en el sistema:
    - Windows: descarga desde https://ffmpeg.org/download.html y agrega
      la carpeta 'bin' al PATH (o instala con `winget install ffmpeg`)
    - Verifica con: ffmpeg -version

USO:
    python transcribir_video.py "ruta/al/video.mp4"

    Opcional, elegir tamaño de modelo (más grande = más preciso, más lento):
    python transcribir_video.py "video.mp4" --modelo medium

    Tamaños disponibles: tiny, base, small, medium, large-v3
    Para español, "medium" da muy buen balance precisión/velocidad en CPU.
    Si tienes GPU NVIDIA con CUDA, usa large-v3 y pasa --device cuda.

SALIDA:
    Genera en la misma carpeta del video:
    - <nombre>_transcripcion.docx   (documento Word, con y sin timestamps)
    - <nombre>_transcripcion.txt    (texto plano, por si lo necesitas)
"""

import argparse
import subprocess
import sys
from pathlib import Path


def extraer_audio(video_path: Path, audio_path: Path):
    """Extrae el audio del video a WAV mono 16kHz (formato ideal para Whisper)."""
    print(f"[1/3] Extrayendo audio de: {video_path.name}")
    cmd = [
        "ffmpeg", "-y", "-i", str(video_path),
        "-vn",  # sin video
        "-ac", "1",  # mono
        "-ar", "16000",  # 16kHz
        "-acodec", "pcm_s16le",
        str(audio_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print("Error extrayendo audio con ffmpeg:")
        print(result.stderr[-2000:])
        sys.exit(1)
    print(f"    Audio extraído en: {audio_path}")


def transcribir(audio_path: Path, modelo: str, device: str, idioma: str | None):
    """Transcribe el audio usando faster-whisper. Devuelve lista de segmentos."""
    from faster_whisper import WhisperModel

    print(f"[2/3] Cargando modelo '{modelo}' (device={device})... "
          f"la primera vez puede tardar por la descarga.")
    compute_type = "float16" if device == "cuda" else "int8"
    model = WhisperModel(modelo, device=device, compute_type=compute_type)

    print("    Transcribiendo... (esto puede tardar varios minutos en videos largos)")
    segments_gen, info = model.transcribe(
        str(audio_path),
        language=idioma,          # None = detección automática
        vad_filter=True,          # ignora silencios largos, clave para videos de 1h+
        vad_parameters=dict(min_silence_duration_ms=500),
        beam_size=5,
    )

    print(f"    Idioma detectado: {info.language} "
          f"(probabilidad {info.language_probability:.2f})")

    segments = []
    for seg in segments_gen:
        segments.append(seg)
        # Progreso simple en consola
        mins, secs = divmod(int(seg.end), 60)
        hh, mins = divmod(mins, 60)
        print(f"    [{hh:02d}:{mins:02d}:{secs:02d}] {seg.text.strip()}")

    return segments


def formatear_timestamp(segundos: float) -> str:
    total = int(segundos)
    hh, resto = divmod(total, 3600)
    mm, ss = divmod(resto, 60)
    return f"{hh:02d}:{mm:02d}:{ss:02d}"


def generar_documento(segments, salida_docx: Path, salida_txt: Path, nombre_video: str):
    """Genera un .docx con la transcripción (con timestamps) y un .txt plano."""
    from docx import Document
    from docx.shared import Pt

    print(f"[3/3] Generando documento: {salida_docx.name}")

    # --- TXT plano (solo el texto corrido, útil para copiar/pegar) ---
    texto_plano = " ".join(seg.text.strip() for seg in segments)
    salida_txt.write_text(texto_plano, encoding="utf-8")

    # --- DOCX con formato ---
    doc = Document()
    doc.add_heading(f"Transcripción: {nombre_video}", level=1)

    doc.add_heading("Texto completo", level=2)
    p = doc.add_paragraph(texto_plano)
    p.style.font.size = Pt(11)

    doc.add_page_break()
    doc.add_heading("Transcripción con marcas de tiempo", level=2)
    for seg in segments:
        ts = formatear_timestamp(seg.start)
        para = doc.add_paragraph()
        run_ts = para.add_run(f"[{ts}] ")
        run_ts.bold = True
        para.add_run(seg.text.strip())

    doc.save(salida_docx)
    print(f"    Documento guardado en: {salida_docx}")
    print(f"    Texto plano guardado en: {salida_txt}")


def main():
    parser = argparse.ArgumentParser(description="Transcribe un video largo a texto (gratis, local).")
    parser.add_argument("video", type=str, help="Ruta al archivo de video (mp4, mkv, mov, etc.)")
    parser.add_argument("--modelo", type=str, default="medium",
                         choices=["tiny", "base", "small", "medium", "large-v3"],
                         help="Tamaño del modelo Whisper (default: medium)")
    parser.add_argument("--device", type=str, default="cpu", choices=["cpu", "cuda"],
                         help="Usa 'cuda' si tienes GPU NVIDIA compatible (default: cpu)")
    parser.add_argument("--idioma", type=str, default="es",
                         help="Código de idioma (ej: es, en). Usa 'auto' para detección automática.")
    args = parser.parse_args()

    video_path = Path(args.video).expanduser().resolve()
    if not video_path.exists():
        print(f"No se encontró el archivo: {video_path}")
        sys.exit(1)

    idioma = None if args.idioma == "auto" else args.idioma

    audio_path = video_path.with_suffix(".wav")
    salida_docx = video_path.with_name(video_path.stem + "_transcripcion.docx")
    salida_txt = video_path.with_name(video_path.stem + "_transcripcion.txt")

    extraer_audio(video_path, audio_path)
    segments = transcribir(audio_path, args.modelo, args.device, idioma)
    generar_documento(segments, salida_docx, salida_txt, video_path.stem)

    # Limpieza del wav temporal
    audio_path.unlink(missing_ok=True)

    print("\n✅ Listo.")


if __name__ == "__main__":
    main()
